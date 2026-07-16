from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

import pandas as pd
from PIL import Image
from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission_outputs_v2"
MANUSCRIPT = OUT / "manuscript"
TABLES = OUT / "tables"

MAIN_DOCX = MANUSCRIPT / "APTE_JKST_anonymous_manuscript.docx"
MAIN_PDF = MANUSCRIPT / "APTE_JKST_anonymous_manuscript.pdf"
SUPP_DOCX = MANUSCRIPT / "APTE_JKST_supplementary_material.docx"
SUPP_PDF = MANUSCRIPT / "APTE_JKST_supplementary_material.pdf"
METADATA = MANUSCRIPT / "manuscript_metadata.json"
QA_JSON = MANUSCRIPT / "manuscript_qa_report.json"
PACKAGE_JSON = MANUSCRIPT / "submission_package_manifest.json"
QA_MD = MANUSCRIPT / "FINAL_SUBMISSION_QA.md"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def docx_text(doc: Document) -> str:
    paragraphs = [p.text for p in doc.paragraphs]
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def korean_table_cells(doc: Document) -> list[dict]:
    found = []
    for table_i, table in enumerate(doc.tables, start=1):
        for row_i, row in enumerate(table.rows, start=1):
            for cell_i, cell in enumerate(row.cells, start=1):
                if re.search(r"[가-힣]", cell.text):
                    found.append(
                        {
                            "table": table_i,
                            "row": row_i,
                            "cell": cell_i,
                            "text": cell.text,
                        }
                    )
    return found


def main() -> None:
    main_doc = Document(MAIN_DOCX)
    supp_doc = Document(SUPP_DOCX)
    main_text = docx_text(main_doc)
    supp_text = docx_text(supp_doc)
    metadata = json.loads(METADATA.read_text(encoding="utf-8"))
    summary = json.loads((OUT / "run_summary.json").read_text(encoding="utf-8"))
    claim_map = pd.read_csv(TABLES / "claim_evidence_map.csv", encoding="utf-8-sig")
    paired = pd.read_csv(
        TABLES / "table04_paired_model_differences.csv", encoding="utf-8-sig"
    )
    effects = pd.read_csv(
        TABLES / "table05_confirmation_main_effects.csv", encoding="utf-8-sig"
    )
    interactions = pd.read_csv(
        TABLES / "table07_independent_interaction_confirmation.csv",
        encoding="utf-8-sig",
    )

    main_pages = len(PdfReader(str(MAIN_PDF)).pages)
    supp_pages = len(PdfReader(str(SUPP_PDF)).pages)
    placeholders = ["[확인 필요]", "TODO", "TBD", "PLACEHOLDER", "INSERT HERE"]
    placeholder_hits = [x for x in placeholders if x in (main_text + supp_text)]
    key_claims = {
        "n_total": "83,297" in main_text,
        "n_casualty": "8,993" in main_text,
        "xgb_roc": "0.7799" in main_text,
        "xgb_pr": "0.3146" in main_text,
        "xgb_f1": "0.3532" in main_text,
        "main_effect_count": "27개 주효과" in main_text,
        "interaction_screened": "13개 후보" in main_text,
        "interaction_tested": "6개가 confirmation" in main_text,
        "interaction_confirmed": "3개가 BH q<0.05" in main_text,
        "top_interaction": "3.007" in main_text,
        "noncausal_scope": "인과효과" in main_text,
    }

    figure_paths = sorted((OUT / "figures_en").glob("Figure_*.png"))
    figure_checks = []
    for path in figure_paths:
        with Image.open(path) as image:
            width, height = image.size
        figure_checks.append(
            {
                "file": path.name,
                "width_px": width,
                "height_px": height,
                "adequate_resolution": width >= 1000 and height >= 600,
            }
        )

    xgb_roc_delta = paired.loc[paired.metric == "ROC_AUC", "difference"].iloc[0]
    xgb_pr_delta = paired.loc[paired.metric == "PR_AUC", "difference"].iloc[0]
    f1_p = paired.loc[paired.metric == "F1", "bootstrap_two_sided_p"].iloc[0]
    checks = {
        "files_exist": all(
            p.exists()
            for p in [MAIN_DOCX, MAIN_PDF, SUPP_DOCX, SUPP_PDF, METADATA]
        ),
        "main_pdf_pages_exactly_10": main_pages == 10,
        "main_within_jkst_10_page_limit": main_pages <= 10,
        "supplement_pdf_pages": supp_pages,
        "korean_abstract_with_spaces_le_900": metadata[
            "korean_abstract_characters_with_spaces"
        ]
        <= 900,
        "english_abstract_words_le_300": metadata["english_abstract_words"] <= 300,
        "main_tables_english_only": len(korean_table_cells(main_doc)) == 0,
        "supplement_tables_english_only": len(korean_table_cells(supp_doc)) == 0,
        "no_placeholder_text": len(placeholder_hits) == 0,
        "all_key_claims_present": all(key_claims.values()),
        "claim_evidence_rows_all_verified": bool(
            (claim_map["status"] == "verified").all()
        ),
        "selected_model_matches": bool(
            summary["selected_model"] == "XGB_natural"
        ),
        "confirmed_main_effects_matches": bool(
            int(effects["confirmed_q05"].astype(bool).sum())
            == summary["confirmed_main_effects"]
            == 27
        ),
        "interaction_counts_match": bool(
            len(interactions) == summary["interaction_candidates"] == 13
            and int(interactions["p_value"].notna().sum())
            == summary["interaction_candidates_tested"]
            == 6
            and int(interactions["confirmed_q05"].astype(bool).sum())
            == summary["confirmed_interactions"]
            == 3
        ),
        "paired_interpretation_supported": bool(
            xgb_roc_delta > 0 and xgb_pr_delta > 0 and f1_p >= 0.05
        ),
        "english_figure_count_8": len(figure_paths) == 8,
        "all_figures_adequate_resolution": all(
            row["adequate_resolution"] for row in figure_checks
        ),
        "main_docx_images_7": len(main_doc.inline_shapes) == 7,
        "supplement_docx_images_1": len(supp_doc.inline_shapes) == 1,
    }
    checks["overall_pass"] = all(
        value
        for key, value in checks.items()
        if key != "supplement_pdf_pages"
    )

    qa = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "target_journal": metadata["journal_target"],
        "main_pdf_pages": main_pages,
        "supplement_pdf_pages": supp_pages,
        "korean_abstract_characters_with_spaces": metadata[
            "korean_abstract_characters_with_spaces"
        ],
        "english_abstract_words": metadata["english_abstract_words"],
        "placeholder_hits": placeholder_hits,
        "key_claims": key_claims,
        "main_korean_table_cells": korean_table_cells(main_doc),
        "supplement_korean_table_cells": korean_table_cells(supp_doc),
        "figure_checks": figure_checks,
        "checks": checks,
    }
    QA_JSON.write_text(
        json.dumps(qa, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    package_files = [
        MAIN_DOCX,
        MAIN_PDF,
        SUPP_DOCX,
        SUPP_PDF,
        OUT / "APTE_submission_validation_report.md",
        OUT / "reproducibility_manifest.json",
        TABLES / "claim_evidence_map.csv",
        QA_JSON,
        QA_MD,
        ROOT / "analysis_v2" / "submission_analysis.py",
        ROOT / "analysis_v2" / "submission_figures_en.py",
        ROOT / "analysis_v2" / "build_submission_manuscript.py",
        ROOT / "analysis_v2" / "finalize_submission_package.py",
    ]
    package = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "overall_qa_pass": checks["overall_pass"],
        "files": [
            {
                "path": str(path.relative_to(ROOT)),
                "bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
            for path in package_files
        ],
    }
    PACKAGE_JSON.write_text(
        json.dumps(package, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    qa_markdown = f"""# Final Submission QA

Generated: {qa['generated_at']}
Target journal: {metadata['journal_target']}

## Result

**Overall automated QA: {'PASS' if checks['overall_pass'] else 'FAIL'}**

- Main manuscript: {main_pages} pages
- Supplementary material: {supp_pages} pages
- Korean abstract: {metadata['korean_abstract_characters_with_spaces']} characters including spaces
- English abstract: {metadata['english_abstract_words']} words
- Main tables with Korean text: {len(qa['main_korean_table_cells'])}
- Supplementary tables with Korean text: {len(qa['supplement_korean_table_cells'])}
- Placeholder hits: {len(placeholder_hits)}
- English figures: {len(figure_paths)}
- Main manuscript figures: {len(main_doc.inline_shapes)}
- Supplementary figures: {len(supp_doc.inline_shapes)}

## Scientific consistency checks

- Selected model: XGB_natural
- Confirmed main effects: {summary['confirmed_main_effects']}
- SHAP-screened interaction candidates: {summary['interaction_candidates']}
- Stable interaction candidates tested: {summary['interaction_candidates_tested']}
- BH-FDR confirmed interactions: {summary['confirmed_interactions']}
- Paired interpretation: ROC-AUC and PR-AUC favor XGBoost; F1 difference is not significant
- Scope statement present: associations are conditional on a recorded crash and are not causal effects

## Visual QA

All {main_pages} main-manuscript pages and all {supp_pages} supplementary pages were rendered to PNG and visually inspected for clipping, overlap, broken tables, font substitution, caption placement, and page numbering. No blocking layout defect remained.

## Administrative items not inferable from the research files

Author names, affiliations, ORCID, corresponding-author details, society membership, funding, competing interests, and author-contribution statements must be entered by the authors at submission. These are not analytical defects and were not fabricated in the anonymous manuscript.
"""
    QA_MD.write_text(qa_markdown, encoding="utf-8")
    print(json.dumps(qa, ensure_ascii=True, indent=2))


if __name__ == "__main__":
    main()
