from __future__ import annotations

import json
import math
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission_outputs_v2"
TABLES = OUT / "tables"
FIGURES = OUT / "figures_en"
MANUSCRIPT_DIR = OUT / "manuscript"
MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)

MAIN_DOCX = MANUSCRIPT_DIR / "APTE_JKST_anonymous_manuscript.docx"
SUPP_DOCX = MANUSCRIPT_DIR / "APTE_JKST_supplementary_material.docx"
META_JSON = MANUSCRIPT_DIR / "manuscript_metadata.json"

FONT_KO = "Malgun Gothic"
FONT_EN = "Calibri"
NAVY = "17365D"
BLUE = "2F5597"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
DARK_GRAY = "404040"
WHITE = "FFFFFF"


TITLE_KO = (
    "독립 표본 탐색–확인 설계를 적용한 고속도로 교통사고 인명피해 발생 요인 분석: "
    "XGBoost-SHAP과 강건 로지스틱 회귀의 결합"
)
TITLE_EN = (
    "An Independent-Sample Screening–Confirmation Analysis of Casualty Occurrence "
    "in Expressway Crashes: Integrating XGBoost-SHAP with Robust Logistic Regression"
)

ABSTRACT_KO = (
    "본 연구는 고속도로 교통사고 83,297건에서 인명피해 발생의 예측요인과 비선형 상호작용을 "
    "탐색하되, 탐색과 확인의 표본 중복으로 인한 선택편향을 차단하는 것을 목적으로 한다. 자료를 "
    "학습 50%, 검증 15%, SHAP 탐색 15%, 독립 확인 20%로 층화 분할하고, 불균형 처리·모형선택·"
    "확률보정·분류 임계값 결정을 학습·검증 단계에 한정하였다. 검증 PR-AUC가 가장 높은 자연분포 "
    "XGBoost를 선택한 결과, 독립 확인표본에서 ROC-AUC 0.780, PR-AUC 0.315, F1 0.353을 보였다. "
    "로지스틱 회귀 대비 ROC-AUC와 PR-AUC의 개선은 작지만 유의했으며 F1 차이는 유의하지 않았다. "
    "HC3 강건 로지스틱 회귀에서는 정차차량, 정체, 차로변경, 졸음, 주시태만, 화물·승합차 등이 "
    "높은 조정 오즈와 연관되었다. SHAP으로 선별한 13개 상호작용 후보 중 안정성 기준을 충족한 "
    "6개를 독립 검정하여 승합차×대형 차급 등 3개를 BH-FDR 기준으로 확인하였다. 본 결과는 "
    "사고 발생 조건부 인명피해 연관성이며 인과효과로 해석할 수 없다."
)

ABSTRACT_EN = (
    "This study investigated predictors and nonlinear interactions associated with "
    "casualty occurrence among 83,297 expressway crashes while preventing selection "
    "bias caused by reusing the same observations for exploration and inference. The "
    "data were stratified into training (50%), validation (15%), SHAP screening (15%), "
    "and independent confirmation (20%) sets. Resampling, model selection, Platt "
    "calibration, and threshold selection were restricted to the training and "
    "validation stages. Natural-prevalence XGBoost achieved the highest validation "
    "precision–recall area under the curve (PR-AUC) and yielded a confirmation ROC-AUC "
    "of 0.780, PR-AUC of 0.315, and F1 score of 0.353. Relative to logistic regression, "
    "XGBoost showed small but statistically detectable improvements in ROC-AUC, "
    "PR-AUC, and Brier score, whereas the F1 difference was not significant. HC3-robust "
    "logistic regression identified higher adjusted odds associated with stopped "
    "vehicles, congestion, lane changes, drowsiness, inattention, trucks, and vans. "
    "Of 13 interactions screened using SHAP, six satisfied prespecified confirmation "
    "cell-size rules and three survived Benjamini–Hochberg false-discovery-rate "
    "control, including a van-by-large-vehicle-class interaction. The findings concern "
    "casualty occurrence conditional on a recorded crash and should not be interpreted "
    "as causal effects or population crash risks."
)

KEYWORDS_KO = "고속도로 교통사고; 인명피해; XGBoost; SHAP; 상호작용; 독립 확인표본"
KEYWORDS_EN = (
    "Expressway crash; Casualty occurrence; XGBoost; SHAP; Interaction; "
    "Independent confirmation"
)

VARIABLE_EN = {
    "주사고원인": "Primary cause",
    "교통장애요인": "Traffic obstruction",
    "발생지점": "Crash location",
    "운전자상태": "Driver condition",
    "사고시도로환경": "Road environment",
    "원인차차종": "Causal vehicle",
    "날씨": "Weather",
    "조명시설": "Lighting",
    "노면상태": "Surface condition",
    "연령구분": "Driver age",
    "사고직전차량조작": "Pre-crash maneuver",
    "차종구분": "Vehicle class",
    "절성토구분": "Embankment/cut class",
    "평면선형": "Horizontal alignment",
    "주말여부": "Day type",
    "포장구분": "Pavement type",
    "성별": "Sex",
}

CATEGORY_EN = {
    "20세미만": "Under 20",
    "20대": "20s",
    "30대": "30s",
    "40대": "40s",
    "50대": "50s",
    "60세이상": "60 or older",
    "SUV형": "SUV",
    "건조": "Dry",
    "경형": "Mini",
    "과속": "Speeding",
    "기타": "Other",
    "낙하물": "Fallen object",
    "남": "Male",
    "노면잡물": "Road debris",
    "눈": "Snow",
    "단독차량화재": "Single-vehicle fire",
    "대형": "Large",
    "도로사정": "Road condition",
    "동물침입": "Animal intrusion",
    "램프": "Ramp",
    "맑음": "Clear",
    "미끄러운노면": "Slippery surface",
    "미상": "Unknown",
    "미작동": "Not operating",
    "본선": "Mainline",
    "비": "Rain",
    "성토고 15M 이상": "Embankment ≥15 m",
    "성토고 2M 미만": "Embankment <2 m",
    "성토고 2M 이상 5M 미만": "Embankment 2–<5 m",
    "소형": "Small",
    "습기": "Damp",
    "승용": "Passenger car",
    "승합": "Van",
    "시거장애": "Visibility obstruction",
    "아스팔트": "Asphalt",
    "안전거리미확보": "Insufficient headway",
    "없음": "None",
    "여": "Female",
    "영업소(TG)": "Tollgate (TG)",
    "우커브": "Right curve",
    "운행차로주행": "Lane keeping",
    "음주": "Alcohol",
    "일반화물": "General freight",
    "작동": "Operating",
    "장애없음": "No obstruction",
    "적설": "Snow-covered",
    "적재불량": "Improper loading",
    "절토부 10M 미만": "Cut <10 m",
    "절토부 10M 이상": "Cut ≥10 m",
    "정상": "Normal",
    "정차차량": "Stopped vehicle",
    "정체": "Congestion",
    "졸음": "Drowsiness",
    "좌커브": "Left curve",
    "주말": "Weekend",
    "주시태만": "Inattention",
    "중형": "Medium",
    "직선": "Straight",
    "차량결함기타": "Other vehicle defect",
    "차량기타": "Other vehicle factor",
    "차로변경": "Lane change",
    "추월불량": "Unsafe passing",
    "콘크리트": "Concrete",
    "타이어파손": "Tire failure",
    "탑차": "Box truck",
    "터널": "Tunnel",
    "트레일러": "Trailer",
    "특수": "Special vehicle",
    "평일": "Weekday",
    "평지": "Level",
    "포트홀": "Pothole",
    "피로": "Fatigue",
    "해당없음(주간)": "Not applicable (daytime)",
    "핸들과대조작": "Excessive steering",
    "화물": "Truck",
    "휴게소": "Service area",
    "흐림": "Cloudy",
}


def translate_feature(feature: str) -> str:
    if "_" not in feature:
        return CATEGORY_EN.get(feature, VARIABLE_EN.get(feature, feature))
    variable, category = feature.split("_", 1)
    return (
        f"{VARIABLE_EN.get(variable, variable)}="
        f"{CATEGORY_EN.get(category, category)}"
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:{}".format(key)), str(value))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_cell_margins(cell, top=35, start=45, bottom=35, end=45) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size: float, bold: bool = False, color: str = DARK_GRAY,
                 italic: bool = False, font: str = FONT_KO) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, 8, color="666666", font=FONT_EN)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE) -> None:
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_EN)
    r_fonts.set(qn("w:hAnsi"), FONT_EN)
    r_fonts.set(qn("w:eastAsia"), FONT_KO)
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    r_pr.append(sz)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_document(doc: Document, running_title: str) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_KO
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.15
    pf.space_after = Pt(3)
    pf.widow_control = True

    for style_name, size, color, before, after in (
        ("Title", 16.5, NAVY, 0, 5),
        ("Subtitle", 11.3, DARK_GRAY, 0, 8),
        ("Heading 1", 12.0, NAVY, 9, 3),
        ("Heading 2", 10.2, BLUE, 6, 2),
        ("Heading 3", 9.4, DARK_GRAY, 4, 1),
    ):
        style = styles[style_name]
        style.font.name = FONT_KO
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Table Text" not in styles:
        table_style = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_style = styles["Table Text"]
    table_style.font.name = FONT_EN
    table_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    table_style.font.size = Pt(7.4)
    table_style.paragraph_format.space_after = Pt(0)
    table_style.paragraph_format.line_spacing = 1.0

    if "Caption APTE" not in styles:
        cap = styles.add_style("Caption APTE", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Caption APTE"]
    cap.font.name = FONT_EN
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    cap.font.size = Pt(8.2)
    cap.font.bold = True
    cap.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(2)
    cap.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(running_title)
    set_run_font(r, 7.5, color="777777", font=FONT_EN)
    footer = section.footer
    add_page_number(footer.paragraphs[0])

    doc.core_properties.title = TITLE_EN
    doc.core_properties.subject = "Anonymous manuscript for academic peer review"
    doc.core_properties.keywords = KEYWORDS_EN
    doc.core_properties.comments = (
        "academic_manuscript_jkst named override of narrative_proposal preset: "
        "A4, 16 mm vertical and 18 mm horizontal margins, Malgun Gothic 9.2 pt, "
        "1.15 line spacing, English tables/figures/references."
    )


def add_text(doc: Document, text: str, bold_lead: str | None = None,
             italic: bool = False, align=None, after: float | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, 9.2, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, 9.2, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, 9.2, italic=italic)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Mm(5)
        p.paragraph_format.first_line_indent = Mm(-2.5)
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(item)
        set_run_font(r, 9.0)


def add_caption(doc: Document, label: str, text: str, above: bool = True) -> None:
    p = doc.add_paragraph(style="Caption APTE")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if above else WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}. {text}")
    set_run_font(r, 8.2, bold=True, font=FONT_EN)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_cm: list[float] | None = None,
    font_size: float = 7.2,
    header_font_size: float = 7.2,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for j, header in enumerate(headers):
        cell = hdr.cells[j]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, NAVY)
        set_table_cell_margins(cell)
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Text"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, header_font_size, bold=True, color=WHITE, font=FONT_EN)
        if widths_cm:
            cell.width = Cm(widths_cm[j])
    for i, row in enumerate(rows):
        tr = table.add_row()
        prevent_row_split(tr)
        for j, value in enumerate(row):
            cell = tr.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_table_cell_margins(cell)
            if i % 2 == 1:
                set_cell_shading(cell, "F7F9FC")
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            r = p.add_run(str(value))
            font = FONT_KO if re.search(r"[가-힣]", str(value)) else FONT_EN
            set_run_font(r, font_size, font=font)
            if widths_cm:
                cell.width = Cm(widths_cm[j])
    edge = {"val": "single", "sz": "4", "color": "A6A6A6"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top=edge,
                bottom=edge,
                left=edge,
                right=edge,
            )
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(1)


def add_table_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_together = True
    r1 = p.add_run("Note. ")
    set_run_font(r1, 7.5, bold=True, font=FONT_EN)
    r2 = p.add_run(text)
    set_run_font(r2, 7.5, font=FONT_EN)


def add_figure(doc: Document, path: Path, label: str, caption: str,
               width_in: float = 6.45) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    add_caption(doc, label, caption, above=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def fmt(x, digits=3) -> str:
    if pd.isna(x):
        return "—"
    return f"{float(x):.{digits}f}"


def fmt_p(x) -> str:
    if pd.isna(x):
        return "—"
    value = float(x)
    if value < 0.001:
        return "<0.001"
    return f"{value:.3f}"


def clean_feature(text: str) -> str:
    compact = {
        "원인차차종_승합": "Van",
        "차종구분_대형": "Large class",
        "주사고원인_주시태만": "Inattention",
        "발생지점_영업소(TG)": "Tollgate",
        "차종구분_소형": "Small class",
        "발생지점_램프": "Ramp",
        "사고직전차량조작_핸들과대조작": "Excessive steering",
        "주사고원인_졸음": "Drowsiness",
        "조명시설_없음": "No lighting",
    }
    return compact.get(text, translate_feature(text))


def build_main_document() -> dict:
    split = pd.read_csv(TABLES / "table01_split_summary.csv", encoding="utf-8-sig")
    validation = pd.read_csv(
        TABLES / "table02_validation_model_selection.csv", encoding="utf-8-sig"
    )
    performance = pd.read_csv(
        TABLES / "table03_confirmation_performance_ci.csv", encoding="utf-8-sig"
    )
    paired = pd.read_csv(
        TABLES / "table04_paired_model_differences.csv", encoding="utf-8-sig"
    )
    effects = pd.read_csv(
        TABLES / "table05_confirmation_main_effects.csv", encoding="utf-8-sig"
    )
    interactions = pd.read_csv(
        TABLES / "table07_independent_interaction_confirmation.csv",
        encoding="utf-8-sig",
    )

    doc = Document()
    configure_document(doc, "Anonymous manuscript | APTE independent confirmation study")

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE_KO)
    set_run_font(r, 16.5, bold=True, color=NAVY)
    p2 = doc.add_paragraph(style="Subtitle")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(TITLE_EN)
    set_run_font(r, 11.3, bold=True, color=DARK_GRAY, font=FONT_EN)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("Research Article · Anonymous manuscript for peer review")
    set_run_font(r, 8.4, italic=True, color="666666", font=FONT_EN)

    doc.add_heading("Abstract", level=1)
    add_text(doc, ABSTRACT_EN)
    p = doc.add_paragraph()
    r1 = p.add_run("Keywords: ")
    set_run_font(r1, 8.8, bold=True, font=FONT_EN)
    r2 = p.add_run(KEYWORDS_EN)
    set_run_font(r2, 8.8, font=FONT_EN)

    doc.add_heading("초록", level=1)
    add_text(doc, ABSTRACT_KO)
    p = doc.add_paragraph()
    r1 = p.add_run("주요어: ")
    set_run_font(r1, 8.8, bold=True)
    r2 = p.add_run(KEYWORDS_KO)
    set_run_font(r2, 8.8)

    doc.add_heading("I. 서론", level=1)
    add_text(
        doc,
        "고속도로 교통사고의 인명피해는 사고 발생 여부와 별개로 충돌 당시의 운전자 행동, 차량 "
        "특성, 도로·시설 환경, 교통장애 상황이 복합적으로 결합한 결과다. 전통적인 로지스틱 회귀는 "
        "조정 오즈비와 불확실성을 제공하지만 비선형성과 고차 상호작용을 사전에 지정해야 한다. 반면 "
        "XGBoost는 복잡한 비선형 관계를 학습할 수 있고 SHAP은 개별 예측의 변수 기여를 일관된 "
        "척도로 분해할 수 있다(Chen and Guestrin, 2016; Lundberg and Lee, 2017)."
    )
    add_text(
        doc,
        "그러나 동일 표본에서 SHAP으로 후보를 선별하고 다시 회귀계수의 p값을 계산하면 탐색과 "
        "확인이 분리되지 않아 선택편향이 발생한다. 결과변수를 이용한 변수선택을 분할 전에 수행하거나 "
        "테스트표본에서 분류 임계값을 조정하는 관행 역시 외부 성능을 낙관적으로 만든다(Ambroise and "
        "McLachlan, 2002). 특히 인명피해 발생률이 낮은 자료에서는 ROC-AUC만으로 성능을 설명하기보다 "
        "PR-AUC, 확률보정, 임계값 의존 지표를 함께 제시해야 한다(Saito and Rehmsmeier, 2015)."
    )
    add_text(
        doc,
        "국내 교통안전 연구는 XGBoost, 설명가능 기계학습, 로지스틱·음이항 모형을 활용해 사고 심각도와 "
        "도로환경의 연관성을 분석해 왔다(Yoon et al., 2016; Park et al., 2019; Kwon and Chang, "
        "2021; Moon et al., 2025). 다만 기계학습 기반 후보 탐색과 독립 확인표본의 통계적 검정을 "
        "엄격히 연결한 고속도로 인명피해 연구는 제한적이다."
    )
    add_text(doc, "본 연구의 목적은 다음 세 가지다.")
    add_bullets(
        doc,
        [
            "고속도로 사고 83,297건에서 자연 인명피해 유병률을 보존한 독립 확인 성능을 산출한다.",
            "XGBoost-SHAP 탐색과 HC3 강건 로지스틱 확인을 독립 표본으로 분리한다.",
            "재현 가능한 주효과와 상호작용만을 교통안전 관리 관점에서 해석하고 한계를 명시한다.",
        ],
    )

    doc.add_heading("II. 선행연구와 연구 공백", level=1)
    add_text(
        doc,
        "사고 심각도 연구에서 로지스틱 계열 모형은 계수의 방향과 기준범주 대비 효과를 명시하는 "
        "장점이 있다. 국내 연구에서는 작업구간, 기상과 도로기하, 이륜차 사고를 대상으로 사고 "
        "심각도 영향요인이 보고되었다(Yoon et al., 2016; Park et al., 2019; Kwon and Chang, "
        "2021). 최근에는 SHAP과 기계학습을 이용해 비선형 관계와 상호작용을 해석하는 연구가 "
        "확대되고 있다(Moon et al., 2025; Moon and Lee, 2025)."
    )
    add_text(
        doc,
        "설명가능 기계학습의 중요도는 예측모형 내부의 기여를 나타내며, 동일한 기준범주를 둔 조정 "
        "오즈비와 같은 양이 아니다. 따라서 SHAP의 크기나 부호를 곧바로 인과적 위험요인으로 읽지 "
        "않고, 독립 자료에서 사전 안정성 규칙을 충족한 후보만 검정해야 한다. 본 연구는 네 개의 "
        "표본 역할을 분리하고, 후보 생성과 가설검정 사이의 정보 흐름을 차단하는 설계를 적용했다."
    )

    doc.add_heading("III. 자료 및 방법", level=1)
    doc.add_heading("1. 자료와 분석대상", level=2)
    add_text(
        doc,
        "분석자료는 한국도로공사 고속도로 교통사고 자료에서 구축한 사고 단위 자료 83,297건이다. "
        "사망·중상·경상 플래그 중 하나 이상이 1이면 인명피해 발생으로 정의했으며, 양성 사례는 "
        "8,993건(10.796%)이다. 실제 사상자 수는 공개 저장소의 원자료에 존재하지 않아 분석하지 "
        "않았다. 따라서 연구의 추론대상은 사고발생 자체가 아니라 기록된 사고 중 인명피해 동반 "
        "여부다."
    )
    add_text(
        doc,
        "설명변수는 주사고원인, 교통장애요인, 발생지점, 운전자상태, 사고시 도로환경, 원인차차종, "
        "날씨, 조명시설, 노면상태, 연령구분, 사고직전 차량조작, 차종구분, 절성토구분, 평면선형, "
        "주말여부, 포장구분, 성별의 17개 범주형 변수다. 주야구분은 조명시설의 주간 범주와 구조적으로 "
        "중복되어 결과변수와 무관한 사전 규칙으로 제외했다."
    )

    doc.add_heading("2. 독립 탐색–확인 설계", level=2)
    add_text(
        doc,
        "무작위 시드 20260716을 고정하고 결과변수로 층화하여 학습 50%, 검증 15%, SHAP 탐색 "
        "15%, 독립 확인 20%로 분할했다. 범주 인코딩과 불균형 처리는 학습표본에서만 적합했다. "
        "검증표본은 모형 선택, Platt 확률보정, F1 최대 임계값 선택에만 사용했고, SHAP 탐색표본은 "
        "변수 중요도와 상호작용 후보 생성에만 사용했다. 확인표본은 최종 성능, 주효과, 사전 선별 "
        "상호작용 검정에만 한 번 사용했다."
    )
    add_caption(doc, "Table 1", "Independent data roles and outcome prevalence.")
    order = ["train", "validation", "screening", "confirmation_test"]
    label = {
        "train": "Training",
        "validation": "Validation",
        "screening": "SHAP screening",
        "confirmation_test": "Independent confirmation",
    }
    split_rows = []
    for name in order:
        row = split.loc[split["split"] == name].iloc[0]
        split_rows.append(
            [
                label[name],
                f"{float(row['proportion']) * 100:.1f}%",
                f"{int(row['n']):,}",
                f"{int(row['casualty_n']):,}",
                f"{float(row['prevalence']) * 100:.2f}%",
            ]
        )
    add_table(
        doc,
        ["Role", "Share", "Crashes", "Casualty crashes", "Prevalence"],
        split_rows,
        widths_cm=[5.0, 2.2, 2.6, 3.1, 2.8],
    )
    add_table_note(
        doc,
        "All splits were stratified by casualty occurrence. No confirmation outcomes "
        "were used for preprocessing, model selection, calibration, or threshold tuning.",
    )
    add_figure(
        doc,
        FIGURES / "Figure_1_independent_design.png",
        "Figure 1",
        "Independent screening–confirmation design and permitted information flow.",
        width_in=6.45,
    )

    doc.add_heading("3. 예측모형, 불균형 처리와 성능평가", level=2)
    add_text(
        doc,
        "비교모형은 로지스틱 회귀와 자연분포 XGBoost, class-weighted XGBoost, 무작위 과소표집 "
        "XGBoost, SMOTEN XGBoost였다. 표본 재구성은 학습자료에만 적용했다(Chawla et al., 2002; "
        "Lemaître et al., 2017). 양성 비율이 10.8%인 불균형 자료이므로 validation PR-AUC를 "
        "1차 선택지표로 사전 지정했다. 확인표본에서 ROC-AUC, PR-AUC, accuracy, balanced accuracy, "
        "precision, recall, specificity, F1, Brier score와 calibration intercept·slope를 보고했다."
    )
    add_text(
        doc,
        "불확실성은 결과변수를 보존한 1,000회 층화 bootstrap 백분위 구간으로 계산했다. 선택된 "
        "XGBoost와 로지스틱 회귀의 차이는 동일 재표집에서 paired bootstrap으로 평가했다. 0으로 "
        "계산된 bootstrap p값은 유한 반복수의 한계를 반영해 p<0.001로 보고했다."
    )

    doc.add_heading("4. SHAP 탐색과 독립 통계확인", level=2)
    add_text(
        doc,
        "TreeSHAP은 screening 표본에서만 계산했다(Lundberg et al., 2020). 상호작용은 평균 절대 "
        "SHAP interaction과 동시발생당 기여를 함께 고려해 후보를 생성했고 `미상`, `기타`, `-` "
        "범주는 후보에서 제외했다. screening 동시발생 25건 이상을 후보 규칙으로 두었다."
    )
    add_text(
        doc,
        "독립 confirmation 표본에서 모든 주효과를 포함한 이항 GLM을 적합하고 HC3 강건 표준오차를 "
        "사용했다. 주효과 계수 전체에 Benjamini–Hochberg FDR을 적용했다(Benjamini and Hochberg, "
        "1995). 상호작용 후보는 confirmation 동시발생 100건 이상, 사건 10건 이상, 비사건 10건 "
        "이상을 모두 만족할 때만 검정했다. 검정 가능한 전체 후보에 다시 BH-FDR을 적용했다."
    )

    doc.add_heading("IV. 결과", level=1)
    doc.add_heading("1. 모형 선택과 독립 확인 성능", level=2)
    validation_rows = []
    for _, row in validation.iterrows():
        validation_rows.append(
            [
                row["model"],
                fmt(row["ROC_AUC"]),
                fmt(row["PR_AUC"]),
                fmt(row["Recall"]),
                fmt(row["F1"]),
                fmt(row["Brier"]),
            ]
        )
    add_caption(doc, "Table 2", "Validation-set model selection under natural prevalence.")
    add_table(
        doc,
        ["Model", "ROC-AUC", "PR-AUC", "Recall", "F1", "Brier"],
        validation_rows,
        widths_cm=[4.0, 2.2, 2.2, 2.2, 2.0, 2.2],
    )
    add_table_note(
        doc,
        "The primary selection metric was validation PR-AUC. Threshold-dependent "
        "metrics use validation-selected thresholds. RUS = random undersampling.",
    )
    add_text(
        doc,
        "자연분포 XGBoost의 validation PR-AUC는 0.3108로 가장 높아 최종 모형으로 선택되었다. "
        "가중치, 과소표집, SMOTEN은 자연분포 학습보다 PR-AUC를 개선하지 못했다."
    )

    metrics = [
        "ROC_AUC",
        "PR_AUC",
        "Accuracy",
        "Balanced_Accuracy",
        "Precision",
        "Recall",
        "Specificity",
        "F1",
        "Brier",
    ]
    metric_labels = {
        "ROC_AUC": "ROC-AUC",
        "PR_AUC": "PR-AUC",
        "Accuracy": "Accuracy",
        "Balanced_Accuracy": "Balanced accuracy",
        "Precision": "Precision",
        "Recall": "Recall",
        "Specificity": "Specificity",
        "F1": "F1",
        "Brier": "Brier score",
    }
    perf_rows = []
    for m in metrics:
        xgb = performance[(performance.model == "XGB_natural") & (performance.metric == m)].iloc[0]
        logit = performance[(performance.model == "Logistic") & (performance.metric == m)].iloc[0]
        perf_rows.append(
            [
                metric_labels[m],
                f"{fmt(xgb['estimate'])} ({fmt(xgb['ci_low'])}–{fmt(xgb['ci_high'])})",
                f"{fmt(logit['estimate'])} ({fmt(logit['ci_low'])}–{fmt(logit['ci_high'])})",
            ]
        )
    add_caption(doc, "Table 3", "Final performance on the untouched confirmation set.")
    add_table(
        doc,
        ["Metric", "XGBoost (95% CI)", "Logistic (95% CI)"],
        perf_rows,
        widths_cm=[4.4, 5.4, 5.4],
    )
    add_table_note(
        doc,
        "Confidence intervals are from 1,000 stratified bootstrap samples. "
        "Validation-selected thresholds were 0.1528 for XGBoost and 0.1626 for logistic regression.",
    )
    add_text(
        doc,
        "확인표본에서 XGBoost의 ROC-AUC는 0.7799(95% CI 0.7694–0.7899), PR-AUC는 "
        "0.3146(0.2953–0.3350), F1은 0.3532(0.3401–0.3659), Brier score는 "
        "0.0851(0.0839–0.0863)이었다. Calibration intercept 0.0006과 slope 1.0038은 "
        "확인표본의 확률보정이 이상적 기준에 가까웠음을 나타냈다."
    )
    roc = paired[paired.metric == "ROC_AUC"].iloc[0]
    pr = paired[paired.metric == "PR_AUC"].iloc[0]
    f1 = paired[paired.metric == "F1"].iloc[0]
    brier = paired[paired.metric == "Brier"].iloc[0]
    add_text(
        doc,
        f"Paired bootstrap에서 XGBoost–로지스틱 차이는 ROC-AUC +{roc['difference']:.4f}"
        f"({roc['ci_low']:.4f}–{roc['ci_high']:.4f}), PR-AUC +{pr['difference']:.4f}"
        f"({pr['ci_low']:.4f}–{pr['ci_high']:.4f}), Brier score {brier['difference']:.5f}"
        f"({brier['ci_low']:.5f}–{brier['ci_high']:.5f})였고 모두 p<0.001이었다. "
        f"F1 차이는 {f1['difference']:.4f}({f1['ci_low']:.4f}–{f1['ci_high']:.4f}, "
        f"p={f1['bootstrap_two_sided_p']:.3f})로 유의하지 않았다. 즉, XGBoost의 순위 판별과 "
        "확률오차 개선은 작지만 재현되었고 임계값 기반 F1 우위는 확인되지 않았다."
    )
    add_figure(
        doc,
        FIGURES / "Figure_2_confirmation_roc_pr.png",
        "Figure 2",
        "Discrimination on the untouched confirmation set.",
        width_in=6.35,
    )
    add_figure(
        doc,
        FIGURES / "Figure_3_confirmation_calibration.png",
        "Figure 3",
        "Probability calibration on the untouched confirmation set.",
        width_in=4.95,
    )

    doc.add_heading("2. SHAP 탐색과 주효과의 독립 확인", level=2)
    add_text(
        doc,
        "Screening 표본의 평균 절대 SHAP에서 발생지점 본선, 노면잡물, 승용차, 교통장애 없음, "
        "주시태만, 영업소, 졸음이 상위에 위치했다. 이 순위는 예측모형의 기여도이며 기준범주 대비 "
        "조정효과가 아니다."
    )
    add_figure(
        doc,
        FIGURES / "Figure_5_screening_shap.png",
        "Figure 4",
        "SHAP feature importance computed only in the independent screening set.",
        width_in=5.75,
    )

    key_dummies = [
        "교통장애요인=정체",
        "교통장애요인=정차차량",
        "사고직전차량조작=차로변경",
        "주사고원인=졸음",
        "원인차차종=승합",
        "원인차차종=화물",
        "주사고원인=주시태만",
        "주사고원인=안전거리미확보",
        "차종구분=대형",
        "원인차차종=트레일러",
        "발생지점=램프",
        "발생지점=영업소(TG)",
    ]
    contrast_map = {
        "교통장애요인=정체": "Congestion vs no obstruction",
        "교통장애요인=정차차량": "Stopped vehicle vs no obstruction",
        "사고직전차량조작=차로변경": "Lane change vs lane keeping",
        "주사고원인=졸음": "Drowsiness vs speeding",
        "원인차차종=승합": "Van vs passenger car",
        "원인차차종=화물": "Truck vs passenger car",
        "주사고원인=주시태만": "Inattention vs speeding",
        "주사고원인=안전거리미확보": "Insufficient headway vs speeding",
        "차종구분=대형": "Large vs medium class",
        "원인차차종=트레일러": "Trailer vs passenger car",
        "발생지점=램프": "Ramp vs mainline",
        "발생지점=영업소(TG)": "Tollgate vs mainline",
    }
    effect_rows = []
    for dummy in key_dummies:
        row = effects.loc[effects.dummy == dummy].iloc[0]
        effect_rows.append(
            [
                contrast_map[dummy],
                f"{row['OR']:.3f}",
                f"{row['CI_low']:.3f}–{row['CI_high']:.3f}",
                f"{int(row['category_casualty_n'])}/{int(row['category_n'])}",
                f"{row['category_rate'] * 100:.2f}%",
                fmt_p(row["q_value_BH"]),
            ]
        )
    add_caption(doc, "Table 4", "Selected independently confirmed main effects.")
    add_table(
        doc,
        ["Contrast", "Adjusted OR", "95% CI", "Cases/n", "Observed rate", "BH q"],
        effect_rows,
        widths_cm=[6.0, 2.1, 2.7, 2.1, 2.3, 1.8],
        font_size=6.9,
        header_font_size=6.8,
    )
    add_table_note(
        doc,
        "HC3-robust logistic regression was fitted in the independent confirmation set. "
        "Observed rates are descriptive and ORs are mutually adjusted associations. "
        "The complete 27-effect table is provided in Supplementary Table S2.",
    )
    add_text(
        doc,
        "독립 확인표본에서 27개 주효과가 BH q<0.05였다. 정체(OR 3.170), 정차차량(2.893), "
        "차로변경(2.029), 졸음(1.839), 화물차(1.800), 승합차(1.882), 주시태만(1.668)은 "
        "높은 조정 오즈와 연관되었다. 영업소(0.117)와 램프(0.304)는 본선보다 낮은 조정 오즈를 "
        "보였으나 시설의 인과적 보호효과로 해석하지 않았다."
    )
    add_figure(
        doc,
        FIGURES / "Figure_6_main_effects.png",
        "Figure 5",
        "Adjusted odds ratios for independently confirmed main effects.",
        width_in=5.75,
    )

    doc.add_heading("3. 상호작용의 독립 확인", level=2)
    tested = interactions[interactions["p_value"].notna()].copy()
    int_rows = []
    for _, row in tested.iterrows():
        status = "Confirmed" if bool(row["confirmed_q05"]) else "Not confirmed"
        int_rows.append(
            [
                f"{clean_feature(row['feature_1'])} × {clean_feature(row['feature_2'])}",
                f"{row['interaction_OR_ratio']:.3f}",
                f"{row['CI_low']:.3f}–{row['CI_high']:.3f}",
                f"{int(row['confirmation_casualty_n_both'])}/{int(row['confirmation_n_both'])}",
                f"{row['confirmation_rate_both'] * 100:.2f}%",
                fmt_p(row["q_value_BH"]),
                status,
            ]
        )
    add_caption(doc, "Table 5", "Independent confirmation of SHAP-screened interactions.")
    add_table(
        doc,
        ["Interaction", "OR ratio", "95% CI", "Cases/n", "Joint rate", "BH q", "Decision"],
        int_rows,
        widths_cm=[5.0, 1.8, 2.5, 1.8, 1.9, 1.6, 2.3],
        font_size=6.7,
        header_font_size=6.5,
    )
    add_table_note(
        doc,
        "Seven additional screened candidates were excluded by prespecified confirmation "
        "cell-size rules. The FDR family comprised all six testable candidates.",
    )
    add_text(
        doc,
        "13개 후보 중 6개가 confirmation 안정성 기준을 충족했고 3개가 BH q<0.05였다. "
        "승합차×대형 차급의 상호작용 OR ratio는 3.007(95% CI 1.877–4.815)이었다. 둘 다 "
        "아닌 사고의 인명피해율은 10.49%, 승합만 11.11%, 대형만 11.14%였지만 두 조건이 "
        "동시에 존재하면 28.88%(54/187)였다."
    )
    add_text(
        doc,
        "주시태만×영업소의 OR ratio 0.411은 주시태만의 조정 연관성이 본선보다 영업소에서 "
        "약화되었다는 의미이며, 주시태만의 보호효과를 뜻하지 않는다. 영업소×소형 차급의 OR "
        "ratio 2.638은 영업소의 낮은 오즈 연관성이 소형 차급에서 덜 강하다는 의미다. "
        "영업소×트레일러는 동시 셀 인명피해가 4건뿐이어서 검정하지 않았다."
    )
    add_figure(
        doc,
        FIGURES / "Figure_7_interactions.png",
        "Figure 6",
        "Interaction odds-ratio multipliers in the independent confirmation set.",
        width_in=5.9,
    )
    add_figure(
        doc,
        FIGURES / "Figure_8_van_large_rates.png",
        "Figure 7",
        "Observed casualty rates for the van-by-large-vehicle-class interaction.",
        width_in=5.15,
    )

    doc.add_heading("V. 논의", level=1)
    doc.add_heading("1. 예측성능의 의미", level=2)
    add_text(
        doc,
        "XGBoost의 confirmation ROC-AUC 약 0.78은 무작위보다 분명히 높은 중등도 판별력을 의미한다. "
        "그러나 PR-AUC 0.315와 precision 0.256은 양성 유병률 0.108의 불균형 문제와 낮은 "
        "임계값 운용의 대가를 보여준다. Recall 0.570은 인명피해 사고의 약 57%를 포착했지만 "
        "양성 예측 중 실제 인명피해 비율은 약 26%였다. 따라서 이 모형은 자동 의사결정 도구가 "
        "아니라 추가 점검 대상을 우선순위화하는 위험선별 보조도구로 해석해야 한다."
    )
    add_text(
        doc,
        "XGBoost와 로지스틱의 ROC-AUC·PR-AUC 차이는 통계적으로 확인되었지만 절대크기가 작았고 "
        "F1 차이는 없었다. 이는 비선형 모형의 예측 이득을 과장하지 않아야 함을 뜻한다. "
        "로지스틱 회귀는 설명가능성과 성능의 균형이 여전히 경쟁적이며, XGBoost는 확률순위와 "
        "상호작용 후보 탐색에서 보완적 가치를 가진다."
    )

    doc.add_heading("2. 교통안전 함의", level=2)
    add_text(
        doc,
        "정차차량과 정체가 동반된 사고의 높은 오즈는 고속도로에서 속도차가 큰 후미·연쇄충돌의 "
        "위험관리 필요성을 시사한다. 정체 후미 경고, 정차차량 조기탐지, 가변속도제어, 운전자 "
        "즉시경보를 결합한 운영대책을 우선 검토할 수 있다. 차로변경, 졸음, 주시태만, 안전거리 "
        "미확보의 양의 연관성은 운전자상태 모니터링과 차로변경·차간거리 보조기능의 정책적 "
        "중요성과 일치한다."
    )
    add_text(
        doc,
        "화물·승합·트레일러 및 대형 차급의 결과는 차량 질량, 무게중심, 탑승인원, 제동거리와 같은 "
        "구조적 차이를 고려한 차종별 관리가 필요함을 보여준다. 특히 승합×대형 조합의 높은 결합 "
        "인명피해율은 해당 차량군의 운전자 교육, 속도관리, 안전장치 점검, 탑승자 보호정책을 "
        "우선 검토할 근거를 제공한다. 다만 노출량이 없으므로 이 조합의 주행거리당 위험이 높다고 "
        "결론내릴 수는 없다."
    )

    doc.add_heading("3. 방법론적 기여", level=2)
    add_text(
        doc,
        "본 연구의 핵심 기여는 SHAP과 회귀를 단순 병렬 비교하지 않고 역할을 분리한 데 있다. "
        "SHAP은 screening 표본에서 복잡한 후보를 찾고, 회귀는 독립 confirmation 표본에서 "
        "방향·효과크기·불확실성을 평가했다. 후보 안정성 규칙과 전체 검정가능 후보에 대한 FDR을 "
        "사전에 적용해 희소 셀과 사후선택 과장을 줄였다. 이 구조는 설명가능 기계학습의 탐색력을 "
        "유지하면서 통계적 주장의 근거수준을 높이는 실용적 절충안이다."
    )

    doc.add_heading("4. 한계", level=2)
    add_text(
        doc,
        "첫째, 자료는 사고 사례만 포함하므로 일반 운행 중 사고발생 위험을 추정하지 않는다. 둘째, "
        "교통량·주행거리·차량등록대수 등의 노출량이 없어 차종과 시설의 절대위험을 계산할 수 없다. "
        "셋째, 연도와 도로구간 ID가 없어 시간·공간 분할 외부검증과 구간 군집상관 보정을 수행하지 "
        "못했다. 넷째, 세 변수는 원시 사고 ID가 아닌 보존된 행 위치로 복원되어 데이터 계보에 "
        "한계가 있다. 다섯째, 일부 희소범주는 완전분리 또는 낮은 사건 수 때문에 불안정하며, "
        "해당 결과는 부록에 한정했다. 마지막으로 무작위 독립 확인은 내부 재현성을 높이지만 다른 "
        "연도·고속도로망에 대한 외부 타당성을 보장하지 않는다."
    )

    doc.add_heading("VI. 결론", level=1)
    add_text(
        doc,
        "고속도로 교통사고 83,297건을 독립 탐색–확인 구조로 재분석한 결과, 자연분포 XGBoost는 "
        "untouched confirmation 표본에서 ROC-AUC 0.780과 PR-AUC 0.315를 보였다. 로지스틱 "
        "회귀 대비 판별·확률오차의 개선은 작지만 확인되었고 F1 우위는 없었다. 강건 회귀에서는 "
        "정차차량, 정체, 차로변경, 졸음, 주시태만과 상용·대형 차량 특성이 높은 인명피해 오즈와 "
        "연관되었다. 독립 확인된 승합×대형 차급 상호작용은 주효과만으로 포착하기 어려운 우선관리 "
        "조합을 제시했다."
    )
    add_text(
        doc,
        "정책적으로는 정체·정차 조기경보, 운전자 부주의·졸음 관리, 차종별 안전대책을 함께 고려할 "
        "필요가 있다. 학술적으로는 설명가능 기계학습의 후보 탐색과 독립 확인자료의 강건 추론을 "
        "분리함으로써 예측성과 해석가능성 사이의 연결을 강화했다. 후속 연구는 원시 사고 ID, 연도, "
        "도로구간, 교통노출량을 포함해 시간·공간 외부검증과 계층모형을 수행해야 한다."
    )

    doc.add_heading("Data and Code Availability", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "Processed analytical data, split assignments, executable analysis code, "
        "result tables, figures, and a reproducibility manifest are archived in the "
        "accompanying APTE repository. The official source record is available from "
    )
    set_run_font(r, 8.8, font=FONT_EN)
    add_hyperlink(
        p,
        "the Korean Public Data Portal",
        "https://www.data.go.kr/data/100298914/linkedData.do",
    )
    r2 = p.add_run(
        ". The deterministic analysis_id is not the original source-system crash identifier."
    )
    set_run_font(r2, 8.8, font=FONT_EN)

    doc.add_heading("References", level=1)
    references = [
        "Ambroise C. and McLachlan G. J. (2002), Selection bias in gene extraction on the basis of microarray gene-expression data, Proceedings of the National Academy of Sciences, 99(10), 6562–6566. https://doi.org/10.1073/pnas.102102699",
        "Benjamini Y. and Hochberg Y. (1995), Controlling the false discovery rate: A practical and powerful approach to multiple testing, Journal of the Royal Statistical Society: Series B, 57(1), 289–300. https://doi.org/10.1111/j.2517-6161.1995.tb02031.x",
        "Chawla N. V., Bowyer K. W., Hall L. O. and Kegelmeyer W. P. (2002), SMOTE: Synthetic minority over-sampling technique, Journal of Artificial Intelligence Research, 16, 321–357.",
        "Chen T. and Guestrin C. (2016), XGBoost: A scalable tree boosting system, Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785–794. https://doi.org/10.1145/2939672.2939785",
        "Efron B. and Tibshirani R. J. (1993), An Introduction to the Bootstrap, Chapman & Hall/CRC, New York.",
        "Kwon C. W. and Chang H. H. (2021), Comparative analysis of traffic accident severity of two-wheeled vehicles using XGBoost, Journal of Korean Society of Intelligent Transport Systems, 20(4), 1–12. https://doi.org/10.12815/kits.2021.20.4.1",
        "Lemaître G., Nogueira F. and Aridas C. K. (2017), Imbalanced-learn: A Python toolbox to tackle the curse of imbalanced datasets in machine learning, Journal of Machine Learning Research, 18(17), 1–5.",
        "Lundberg S. M., Erion G., Chen H., DeGrave A., Prutkin J. M., Nair B., Katz R., Himmelfarb J., Bansal N. and Lee S. I. (2020), From local explanations to global understanding with explainable AI for trees, Nature Machine Intelligence, 2, 56–67. https://doi.org/10.1038/s42256-019-0138-9",
        "Lundberg S. M. and Lee S. I. (2017), A unified approach to interpreting model predictions, Advances in Neural Information Processing Systems, 30.",
        "Moon J., Jang J. and Lee S. (2025), Analyzing the associations between the fatal accidents of older adult drivers and road environments in Seoul, Korea: Focusing on non-linear relationships and interaction effects using machine learning, Journal of Korean Society of Transportation, 43(2), 161–180. https://doi.org/10.7470/jkst.2025.43.2.161",
        "Moon J. and Lee S. (2025), Analysis of the impact of driver’s visual landscape on traffic accidents in urban functional zones: A hybrid approach combining interpretable machine learning and negative binomial regression, Journal of Korea Planning Association, 60(6), 32–51. https://doi.org/10.17208/jkpa.2025.11.60.6.32",
        "Park S. J., Kho S. Y. and Park H. C. (2019), The effects of road geometry on the injury severity of expressway traffic accident depending on weather conditions, Journal of Korean Society of Intelligent Transport Systems, 18(2), 12–28. https://doi.org/10.12815/kits.2019.18.2.12",
        "Platt J. C. (1999), Probabilistic outputs for support vector machines and comparisons to regularized likelihood methods, Advances in Large Margin Classifiers, MIT Press, 61–74.",
        "Saito T. and Rehmsmeier M. (2015), The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets, PLOS ONE, 10(3), e0118432. https://doi.org/10.1371/journal.pone.0118432",
        "Yoon S. M., Oh C., Park H. J. and Chung B. J. (2016), Identification of factors affecting the crash severity and safety countermeasures toward safer work zone traffic management, Journal of Korean Society of Transportation, 34(4), 354–372. https://doi.org/10.7470/jkst.2016.34.4.354",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(5)
        p.paragraph_format.first_line_indent = Mm(-5)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(ref)
        set_run_font(r, 7.8, font=FONT_EN)

    doc.save(MAIN_DOCX)
    return {
        "main_docx": str(MAIN_DOCX),
        "korean_abstract_characters_no_spaces": len(re.sub(r"\s+", "", ABSTRACT_KO)),
        "korean_abstract_characters_with_spaces": len(ABSTRACT_KO),
        "english_abstract_words": len(re.findall(r"\b[\w–-]+\b", ABSTRACT_EN)),
        "references": len(references),
        "tables": 5,
        "figures": 7,
    }


def build_supplement() -> dict:
    effects = pd.read_csv(
        TABLES / "table05_confirmation_main_effects.csv", encoding="utf-8-sig"
    )
    shap = pd.read_csv(
        TABLES / "table06_screening_shap_importance.csv", encoding="utf-8-sig"
    )
    interactions = pd.read_csv(
        TABLES / "table07_independent_interaction_confirmation.csv",
        encoding="utf-8-sig",
    )
    separated = pd.read_csv(
        TABLES / "confirmation_separated_dummies.csv", encoding="utf-8-sig"
    )
    manifest = json.loads(
        (OUT / "reproducibility_manifest.json").read_text(encoding="utf-8")
    )

    doc = Document()
    configure_document(doc, "Supplementary material | APTE independent confirmation study")
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Supplementary Material")
    set_run_font(r, 17, bold=True, color=NAVY, font=FONT_EN)
    p2 = doc.add_paragraph(style="Subtitle")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(TITLE_EN)
    set_run_font(r, 10.5, bold=True, color=DARK_GRAY, font=FONT_EN)

    doc.add_heading("Supplementary Table S1. Reproducibility specification", level=1)
    repro_rows = [
        ["Random seed", str(manifest["seed"])],
        ["Split", "Training 50%; validation 15%; screening 15%; confirmation 20%"],
        ["Selection metric", manifest["selection_primary_metric"]],
        ["Threshold rule", manifest["threshold_selection"]],
        ["Selected model", manifest["selected_model"]],
        ["Rows / predictors / positive outcomes", "83,297 / 17 / 8,993"],
        ["Python", manifest["versions"]["python"]],
        ["pandas / numpy / scipy", f"{manifest['versions']['pandas']} / {manifest['versions']['numpy']} / {manifest['versions']['scipy']}"],
        ["scikit-learn / XGBoost / SHAP", f"{manifest['versions']['scikit_learn']} / {manifest['versions']['xgboost']} / {manifest['versions']['shap']}"],
        ["statsmodels / imbalanced-learn", f"{manifest['versions']['statsmodels']} / {manifest['versions']['imbalanced_learn']}"],
        ["Processed data SHA-256", manifest["provenance"]["processed_sha256"]],
        ["Partial raw SHA-256", manifest["provenance"]["partial_raw_sha256"]],
    ]
    add_table(doc, ["Item", "Specification"], repro_rows, widths_cm=[5.2, 10.0], font_size=7.0)
    add_table_note(
        doc,
        "analysis_id is deterministic for reproduction and is not the original source-system crash identifier.",
    )

    doc.add_heading("Supplementary Table S2. Complete confirmation main-effect results", level=1)
    effect_rows = []
    for _, row in effects.iterrows():
        effect_rows.append(
            [
                f"{VARIABLE_EN.get(row['variable'], row['variable'])}: "
                f"{CATEGORY_EN.get(row['category'], row['category'])} vs "
                f"{CATEGORY_EN.get(row['reference'], row['reference'])}",
                str(int(row["category_n"])),
                str(int(row["category_casualty_n"])),
                f"{row['category_rate'] * 100:.2f}%",
                f"{row['OR']:.3f}",
                f"{row['CI_low']:.3f}–{row['CI_high']:.3f}",
                fmt_p(row["q_value_BH"]),
                "Yes" if bool(row["confirmed_q05"]) else "No",
            ]
        )
    add_table(
        doc,
        ["Contrast", "n", "Cases", "Rate", "OR", "95% CI", "BH q", "q<0.05"],
        effect_rows,
        widths_cm=[6.0, 1.2, 1.3, 1.5, 1.3, 2.4, 1.4, 1.3],
        font_size=6.2,
        header_font_size=6.0,
    )
    add_table_note(
        doc,
        "HC3-robust standard errors. The table includes all estimable dummy contrasts; "
        "small cells and generic categories should be interpreted cautiously.",
    )

    doc.add_heading("Supplementary Table S3. Complete SHAP screening importance", level=1)
    shap_rows = []
    for _, row in shap.iterrows():
        shap_rows.append(
            [
                translate_feature(row["feature"]),
                VARIABLE_EN.get(row["variable"], row["variable"]),
                f"{row['mean_abs_shap']:.5f}",
                f"{row['mean_shap_when_present']:.5f}",
                str(int(row["screen_n_present"])),
            ]
        )
    add_table(
        doc,
        ["Encoded feature", "Parent variable", "Mean |SHAP|", "Mean SHAP when present", "Screen n"],
        shap_rows,
        widths_cm=[5.5, 3.0, 2.2, 3.0, 1.7],
        font_size=6.3,
        header_font_size=6.2,
    )
    add_table_note(
        doc,
        "SHAP values were computed only in the independent screening set and are model contributions, not causal effects.",
    )

    doc.add_heading("Supplementary Table S4. All screened interaction candidates", level=1)
    interaction_rows = []
    for _, row in interactions.iterrows():
        tested = pd.notna(row["p_value"])
        if tested:
            decision = "Confirmed" if bool(row["confirmed_q05"]) else "Not confirmed"
            or_text = f"{row['interaction_OR_ratio']:.3f}"
            ci_text = f"{row['CI_low']:.3f}–{row['CI_high']:.3f}"
            q_text = fmt_p(row["q_value_BH"])
        else:
            decision = "Excluded by cell rule"
            or_text = "—"
            ci_text = "—"
            q_text = "—"
        interaction_rows.append(
            [
                f"{translate_feature(row['feature_1'])} × "
                f"{translate_feature(row['feature_2'])}",
                str(int(row["screen_sample_n_both"])),
                str(int(row["confirmation_n_both"])),
                str(int(row["confirmation_casualty_n_both"])),
                f"{row['confirmation_rate_both'] * 100:.2f}%",
                or_text,
                ci_text,
                q_text,
                decision,
            ]
        )
    add_table(
        doc,
        ["Candidate", "Screen n", "Confirm n", "Cases", "Joint rate", "OR ratio", "95% CI", "BH q", "Decision"],
        interaction_rows,
        widths_cm=[5.0, 1.3, 1.4, 1.2, 1.5, 1.5, 2.1, 1.2, 2.2],
        font_size=5.9,
        header_font_size=5.7,
    )
    add_table_note(
        doc,
        "Confirmation testing required joint n≥100, events≥10, and non-events≥10. "
        "BH-FDR was applied to all six candidates that satisfied these rules.",
    )

    doc.add_heading("Supplementary Table S5. Complete separation audit", level=1)
    sep_rows = [
        [
            translate_feature(row["dummy"].replace("=", "_", 1)),
            str(int(row["n"])),
            str(int(row["casualty_n"])),
            row["reason"],
        ]
        for _, row in separated.iterrows()
    ]
    add_table(
        doc,
        ["Dummy", "n", "Cases", "Reason"],
        sep_rows,
        widths_cm=[5.0, 1.6, 1.6, 7.0],
        font_size=7.0,
    )
    add_table_note(
        doc,
        "Separated or absent contrasts were removed before estimation and were not treated as confirmed effects.",
    )

    doc.add_heading("Supplementary Figures", level=1)
    add_figure(
        doc,
        FIGURES / "Figure_4_model_metrics.png",
        "Figure S1",
        "Confirmation-set model metrics with 95% bootstrap confidence intervals.",
        width_in=6.35,
    )

    doc.save(SUPP_DOCX)
    return {
        "supplement_docx": str(SUPP_DOCX),
        "supplementary_tables": 5,
        "supplementary_figures": 1,
    }


def main() -> None:
    main_meta = build_main_document()
    supp_meta = build_supplement()
    metadata = {
        **main_meta,
        **supp_meta,
        "title_ko": TITLE_KO,
        "title_en": TITLE_EN,
        "journal_target": "Journal of Korean Society of Transportation",
        "format_profile": {
            "base_preset": "narrative_proposal",
            "named_override": "academic_manuscript_jkst",
            "page": "A4 portrait",
            "margins": "16 mm top/bottom; 18 mm left/right",
            "body": "Malgun Gothic 9.2 pt; 1.15 line spacing",
            "tables_figures_references": "English",
        },
    }
    META_JSON.write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps(metadata, ensure_ascii=True, indent=2))


if __name__ == "__main__":
    main()
